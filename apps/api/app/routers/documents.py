"""Document endpoints — upload, list, detail, delete."""

import csv
import io
import logging
import os
import uuid
from typing import Any

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, status
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.database import get_db
from app.dependencies import get_current_user
from app.models.document import Document
from app.models.extraction import Extraction
from app.models.user import User
from app.schemas.document import (
    DocumentDetailResponse,
    DocumentListResponse,
    DocumentResponse,
    DocumentUploadResponse,
)
from app.tasks.process_document import process_document
from app.utils.pdf_parser import extract_text_from_pdf

router = APIRouter(prefix="/api/documents", tags=["documents"])
logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


class UpdateExtractionFieldRequest(BaseModel):
    field_key: str
    value: Any


# ── POST /api/documents/upload ─────────────────────────


@router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload a PDF contract for processing.

    Saves the file to disk, creates a DB record, and kicks off the
    Celery extraction pipeline asynchronously.
    """
    # Validate file type
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "data": None,
                "error": {"message": "Only PDF files are accepted", "code": "DOC_INVALID_TYPE"},
            },
        )

    # Read and validate file size
    contents = await file.read()
    if len(contents) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail={
                "data": None,
                "error": {"message": "File exceeds 10 MB limit", "code": "DOC_TOO_LARGE"},
            },
        )

    # Build a unique file path: uploads/<team_id>/<uuid>_<filename>
    team_dir = os.path.join(settings.upload_dir, str(user.team_id))
    os.makedirs(team_dir, exist_ok=True)

    file_id = uuid.uuid4()
    safe_name = file.filename or "document.pdf"
    file_path = os.path.join(team_dir, f"{file_id}_{safe_name}")

    # Write file to disk
    with open(file_path, "wb") as f:
        f.write(contents)

    # Pre-extract text from the PDF here in the API process.
    # The API and Celery worker run in separate containers on Railway and do
    # NOT share a filesystem — the worker cannot open a file saved by the API.
    # By extracting text now and storing it in raw_text, the worker can read
    # it from the database and skip the file-read step entirely.
    raw_text: str | None = None
    page_count: int | None = None
    try:
        extracted_text, page_map = extract_text_from_pdf(file_path)
        raw_text = extracted_text
        page_count = len(page_map)
    except Exception as exc:
        logger.warning("Pre-extraction failed, worker will attempt: %s", exc)

    # Create the document record
    document = Document(
        team_id=user.team_id,
        uploaded_by=user.id,
        filename=safe_name,
        file_path=file_path,
        file_size_bytes=len(contents),
        raw_text=raw_text,
        page_count=page_count,
        status="uploaded",
    )
    db.add(document)
    await db.flush()

    # Commit NOW so the Celery worker can see the row immediately.
    # (The get_db dependency will also commit on exit, which is a no-op.)
    await db.commit()

    # Kick off the Celery extraction pipeline
    task = process_document.delay(str(document.id))

    return {
        "data": DocumentUploadResponse(
            document=DocumentResponse.model_validate(document),
            task_id=task.id,
        ).model_dump(),
        "error": None,
    }


# ── GET /api/documents ─────────────────────────────────


@router.get("")
async def list_documents(
    limit: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
    status_filter: str | None = Query(default=None, alias="status"),
    doc_type: str | None = Query(default=None),
    q: str | None = Query(default=None, max_length=500),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """List all documents for the current user's team, newest first.

    Optional filters: status (uploaded|processing|completed|failed),
    doc_type (nda|service_agreement|employment_contract|etc.),
    q (substring match on filename and raw_text content).
    """
    from sqlalchemy import or_

    base = select(Document).where(Document.team_id == user.team_id)

    if status_filter:
        base = base.where(Document.status == status_filter)
    if doc_type:
        base = base.where(Document.doc_type == doc_type)
    if q and q.strip():
        pattern = f"%{q.strip()}%"
        base = base.where(
            or_(
                Document.filename.ilike(pattern),
                Document.raw_text.ilike(pattern),
            )
        )

    # Count total matching
    count_result = await db.execute(select(func.count()).select_from(base.subquery()))
    total = count_result.scalar_one()

    # Fetch page
    result = await db.execute(base.order_by(Document.created_at.desc()).limit(limit).offset(offset))
    documents = result.scalars().all()

    return {
        "data": DocumentListResponse(
            documents=[DocumentResponse.model_validate(d) for d in documents],
            total=total,
            limit=limit,
            offset=offset,
        ).model_dump(),
        "error": None,
    }


# ── GET /api/documents/{id} ───────────────────────────


@router.get("/{document_id}")
async def get_document(
    document_id: uuid.UUID,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get a single document with its extractions, clauses, and annotations."""
    from app.models.annotation import Annotation
    from app.models.clause import Clause
    from app.schemas.document import AnnotationSummary, ClauseResponse

    result = await db.execute(
        select(Document)
        .options(
            selectinload(Document.extractions),
            selectinload(Document.clauses).selectinload(Clause.annotations).selectinload(
                Annotation.user
            ),
        )
        .where(Document.id == document_id, Document.team_id == user.team_id)
    )
    document = result.scalar_one_or_none()

    if document is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "data": None,
                "error": {"message": "Document not found", "code": "DOC_NOT_FOUND"},
            },
        )

    # Build the response, materializing annotations from relationships
    payload = DocumentDetailResponse.model_validate(document).model_dump()
    # Override clauses with annotation-aware serialization
    payload["clauses"] = [
        {
            **ClauseResponse.model_validate(clause).model_dump(),
            "annotations": [
                AnnotationSummary(
                    id=a.id,
                    user_id=a.user_id,
                    user_name=a.user.full_name,
                    content=a.content,
                    created_at=a.created_at,
                ).model_dump()
                for a in clause.annotations
            ],
        }
        for clause in document.clauses
    ]
    return {"data": payload, "error": None}


# ── GET /api/documents/{id}/file ───────────────────────


@router.get("/{document_id}/file")
async def get_document_file(
    document_id: uuid.UUID,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Stream the original uploaded PDF for inline viewing.

    The file lives on the API container's disk. In production that storage
    is ephemeral, so the file may be gone after a restart — callers should
    fall back to the text-based evidence view when this returns 404.
    """
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.team_id == user.team_id,
        )
    )
    document = result.scalar_one_or_none()

    if document is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "data": None,
                "error": {"message": "Document not found", "code": "DOC_NOT_FOUND"},
            },
        )

    if not document.file_path or not os.path.exists(document.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "data": None,
                "error": {
                    "message": "Original file is no longer available on the server.",
                    "code": "DOC_FILE_UNAVAILABLE",
                },
            },
        )

    safe_name = os.path.basename(document.filename).replace('"', "")
    return FileResponse(
        document.file_path,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{safe_name}"'},
    )


# ── POST /api/documents/{id}/reprocess ─────────────────


@router.post("/{document_id}/reprocess")
async def reprocess_document(
    document_id: uuid.UUID,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Re-trigger the extraction pipeline for an existing document.

    Resets the status to 'uploaded' and kicks off a new Celery task.
    Useful when a previous extraction failed or when the pipeline has
    been improved and you want fresh results.
    """
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.team_id == user.team_id,
        )
    )
    document = result.scalar_one_or_none()

    if document is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "data": None,
                "error": {"message": "Document not found", "code": "DOC_NOT_FOUND"},
            },
        )

    if document.status == "processing":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={
                "data": None,
                "error": {
                    "message": "Document is already being processed",
                    "code": "DOC_ALREADY_PROCESSING",
                },
            },
        )

    document.status = "uploaded"
    await db.commit()

    task = process_document.delay(str(document.id))

    return {
        "data": DocumentUploadResponse(
            document=DocumentResponse.model_validate(document),
            task_id=task.id,
        ).model_dump(),
        "error": None,
    }


# ── DELETE /api/documents/{id} ─────────────────────────


@router.delete("/{document_id}")
async def delete_document(
    document_id: uuid.UUID,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete a document, its file on disk, and all related records."""
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.team_id == user.team_id,
        )
    )
    document = result.scalar_one_or_none()

    if document is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "data": None,
                "error": {"message": "Document not found", "code": "DOC_NOT_FOUND"},
            },
        )

    # Delete the file from disk (best-effort)
    if os.path.exists(document.file_path):
        os.remove(document.file_path)

    await db.delete(document)

    return {
        "data": {"message": "Document deleted"},
        "error": None,
    }


# ── PATCH /api/documents/{id}/extractions/{extraction_id} ───
# Inline edit a single extracted field.


@router.patch("/{document_id}/extractions/{extraction_id}")
async def update_extraction_field(
    document_id: uuid.UUID,
    extraction_id: uuid.UUID,
    body: UpdateExtractionFieldRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Update a single field within an extraction's extracted_data JSONB.

    Marks the field as edited (confidence=1.0, source='user') so
    downstream displays know it's a manual correction.
    """
    result = await db.execute(
        select(Extraction)
        .join(Document, Document.id == Extraction.document_id)
        .where(
            Extraction.id == extraction_id,
            Extraction.document_id == document_id,
            Document.team_id == user.team_id,
        )
    )
    extraction = result.scalar_one_or_none()

    if extraction is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "data": None,
                "error": {"message": "Extraction not found", "code": "EXTRACTION_NOT_FOUND"},
            },
        )

    data = dict(extraction.extracted_data or {})
    existing = data.get(body.field_key) or {}
    data[body.field_key] = {
        **existing,
        "value": body.value,
        "confidence": 1.0,
        "source": "user",
    }
    extraction.extracted_data = data

    # SQLAlchemy needs a hint that the JSONB column changed in-place
    from sqlalchemy.orm.attributes import flag_modified

    flag_modified(extraction, "extracted_data")
    await db.commit()
    await db.refresh(extraction)

    return {
        "data": {"extracted_data": extraction.extracted_data},
        "error": None,
    }


# ── GET /api/documents/{id}/export ─────────────────────────


@router.get("/{document_id}/export")
async def export_document(
    document_id: uuid.UUID,
    format: str = Query(default="csv", pattern="^(csv|docx)$"),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Export a document's extraction + clauses as CSV or DOCX."""
    result = await db.execute(
        select(Document)
        .options(
            selectinload(Document.extractions),
            selectinload(Document.clauses),
        )
        .where(
            Document.id == document_id,
            Document.team_id == user.team_id,
        )
    )
    document = result.scalar_one_or_none()

    if document is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={
                "data": None,
                "error": {"message": "Document not found", "code": "DOC_NOT_FOUND"},
            },
        )

    safe_stem = os.path.splitext(document.filename)[0].replace('"', "")

    if format == "csv":
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["Section", "Key", "Value"])
        writer.writerow(["Document", "Filename", document.filename])
        writer.writerow(["Document", "Type", document.doc_type or ""])
        writer.writerow(["Document", "Status", document.status])
        writer.writerow(["Document", "Risk Score", document.risk_score or ""])
        writer.writerow(["Document", "Pages", document.page_count or ""])
        writer.writerow(["Document", "Summary", document.executive_summary or ""])

        for extraction in document.extractions:
            for key, field in (extraction.extracted_data or {}).items():
                value = field.get("value") if isinstance(field, dict) else field
                writer.writerow(["Field", key, "" if value is None else str(value)])

        for clause in document.clauses:
            writer.writerow(
                [
                    f"Clause: {clause.clause_type}",
                    "Risk",
                    clause.risk_level or "",
                ]
            )
            writer.writerow(
                [
                    f"Clause: {clause.clause_type}",
                    "Original",
                    clause.original_text,
                ]
            )
            if clause.plain_summary:
                writer.writerow(
                    [
                        f"Clause: {clause.clause_type}",
                        "Summary",
                        clause.plain_summary,
                    ]
                )
            if clause.suggested_alternative:
                writer.writerow(
                    [
                        f"Clause: {clause.clause_type}",
                        "Suggested",
                        clause.suggested_alternative,
                    ]
                )

        return StreamingResponse(
            iter([buf.getvalue()]),
            media_type="text/csv",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_stem}.csv"',
            },
        )

    # ── DOCX export ──
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail={
                "data": None,
                "error": {
                    "message": "DOCX export unavailable: python-docx not installed.",
                    "code": "EXPORT_DOCX_UNAVAILABLE",
                },
            },
        ) from exc

    docx = DocxDocument()
    docx.add_heading(document.filename, level=1)
    docx.add_paragraph(
        f"Document type: {document.doc_type or 'Pending'}\n"
        f"Status: {document.status}\n"
        f"Risk score: {document.risk_score or 'N/A'}\n"
        f"Pages: {document.page_count or 'N/A'}"
    )

    if document.executive_summary:
        docx.add_heading("Executive Summary", level=2)
        docx.add_paragraph(document.executive_summary)

    if document.extractions:
        docx.add_heading("Extracted Fields", level=2)
        extraction = document.extractions[0]
        for key, field in (extraction.extracted_data or {}).items():
            value = field.get("value") if isinstance(field, dict) else field
            p = docx.add_paragraph()
            run = p.add_run(f"{key.replace('_', ' ').title()}: ")
            run.bold = True
            p.add_run("Not found" if value is None else str(value))

    if document.clauses:
        docx.add_heading("Risk Analysis", level=2)
        for clause in document.clauses:
            docx.add_heading(
                f"{clause.clause_type.replace('_', ' ').title()} "
                f"[{(clause.risk_level or 'unknown').upper()}]",
                level=3,
            )
            quote = docx.add_paragraph()
            run = quote.add_run(f"“{clause.original_text}”")
            run.italic = True
            run.font.size = Pt(10)
            if clause.plain_summary:
                docx.add_paragraph(clause.plain_summary)
            if clause.risk_reason:
                p = docx.add_paragraph()
                p.add_run("Why flagged: ").bold = True
                p.add_run(clause.risk_reason)
            if clause.suggested_alternative:
                p = docx.add_paragraph()
                p.add_run("Suggested: ").bold = True
                p.add_run(clause.suggested_alternative)

    if document.missing_clauses:
        docx.add_heading("Missing Clauses", level=2)
        for clause in document.missing_clauses:
            docx.add_paragraph(str(clause), style="List Bullet")

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": f'attachment; filename="{safe_stem}.docx"',
        },
    )
